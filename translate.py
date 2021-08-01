'''
Author : slayerzeroa

대충 EBP 하기 싫어서 만든 프로그램

아주대학교 금융공학과 후배들에게 널리 배포 쌉가능
'''

#pip install gensim==3.8.3

import tkinter as tk #tkinter 불러오기
import requests
from bs4 import BeautifulSoup as BS
from gensim.summarization.summarizer import summarize
import docx
from datetime import date
import os


root = tk.Tk() # root 지정해주기
root.geometry("400x240") # 박스 만들기

result_address = [] #주소 리스트 생성

def getTextInput(): #입력 불러오는 함수 생성
    result = textExample.get("1.0","end") # 텍스트 위젯에서 첫 문자의 위치는 1.0 텍스트 상자가 끝날 때까지 입력을 읽는다
    result = result.translate(str.maketrans('', '', ' \n\t\r'))
    result_address.append(result)
    root.destroy()

textExample=tk.Text(root, height=10) #textExample의 높이 10
textExample.pack() #pack() => textExample 위젯 위치 배치
btnRead=tk.Button(root, height=1, width=10, text="Start",
                    command=getTextInput) # 버튼 디자인 관련

btnRead.pack() # 버튼 위젯 배치

root.mainloop() # 사용자의 입력을 계속해서 유지해주는 역할

webpage = requests.get(result_address[0]) # web
soup = BS(webpage.content, "html.parser")

data = soup.find_all(class_= "zn-body__paragraph") # CNN 본문내용 html class

data_text_all = "" #빈 문자열

for data_text in data: # 반복
  data_text = data_text.get_text() # html 데이터에서 텍스트만 뽑아오기
  data_text_all = data_text_all + data_text # 문자열 추가

news_summarize = summarize(data_text_all, ratio=0.1, word_count = 130) #요약하는 프로그램


request_url = "https://openapi.naver.com/v1/papago/n2mt" #네이버 papago open api 사용

headers = {"X-Naver-Client-Id": "v_xR89ObiErZTqpWcPmJ", "X-Naver-Client-Secret": "fDcBNe2s5W"} # id, secret id
params = {"source": "en", "target": "ko", "text": news_summarize} # 번역 시작, 끝, 번역할 텍스트
response = requests.post(request_url, headers=headers, data=params) # 정보 넣어주기

result = response.json() #json 파일 저장
translate_result = result['message']['result']['translatedText'] #json 파일 키값 빼오기

path = "./" #폴더 위치
file_list = os.listdir(path) #폴더 내 파일 리스트 가져오기

abs_path = os.getcwd() #폴더 절대 위치

date = date.today() # 오늘 날짜


if '학번_이름_EBP_{}.docx'.format(date) in file_list: #만약 파일리스트에 워드 파일 있으면
    org = docx.Document("{}\\학번_이름_EBP_{}.docx".format(abs_path, date)) # 워드 파일 불러오기
    style = org.styles['Normal']
    style.font.name = 'Calibri'# font명 지정
    style.font.size = docx.shared.Pt(10) # 폰트

    org._body.clear_content() # 기존 파일 내용 삭제
    para_1 = org.add_paragraph() # paragraph 추가
    run_1 = para_1.add_run(translate_result) # news_summarize 내용 추가

    org.save('학번_이름_EBP_{}.docx'.format(date))# 저장해주기

else: # 파일리스트에 워드 파일 없으면
    doc = docx.Document() #doc 지정
    doc.save('학번_이름_EBP_{}.docx'.format(date)) # 빈 워드파일 저장해주기
    style = doc.styles['Normal'] #기본 설정
    style.font.name = 'Calibri'# font명 지정
    style.font.size = docx.shared.Pt(10) # 폰트

    para = doc.add_paragraph() #paragraph 추가
    run = para.add_run(translate_result) # news_summarize 내용 추가

    doc.save('학번_이름_EBP_{}.docx'.format(date)) # 저장해주기